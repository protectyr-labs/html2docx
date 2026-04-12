"""
HTML to DOCX converter with CSS support.

Two-pass algorithm:
  Pass 1: Collect all element IDs for bookmark generation
  Pass 2: Convert HTML elements to Word document objects

Supports: headings, paragraphs, tables, lists, code blocks,
inline formatting (bold/italic/underline/color), CSS variables,
bookmarks, and images.
"""

import re
import os
from dataclasses import dataclass
from typing import Optional, Set, Dict, List
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

__version__ = "0.1.0"

SKIP_TAGS = {"script", "style", "head", "meta", "link", "title", "noscript", "svg"}
CONTAINER_TAGS = {"div", "section", "article", "main", "body", "html", "figure", "header", "footer", "nav", "aside"}
INLINE_TAGS = {"strong", "b", "em", "i", "u", "s", "span", "code", "a", "mark", "small", "sub", "sup"}
HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}

HEADING_SIZES = {"h1": 24, "h2": 20, "h3": 16, "h4": 14, "h5": 12, "h6": 11}
HEADING_LEVELS = {"h1": 0, "h2": 1, "h3": 2, "h4": 3, "h5": 4, "h6": 5}


@dataclass
class RunFormat:
    """Tracks accumulated inline formatting state."""
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: Optional[str] = None  # 6-digit hex without #
    bg_color: Optional[str] = None  # 6-digit hex without #
    font_size: Optional[float] = None
    is_code: bool = False

    def copy(self) -> "RunFormat":
        return RunFormat(
            bold=self.bold, italic=self.italic, underline=self.underline,
            color=self.color, bg_color=self.bg_color,
            font_size=self.font_size, is_code=self.is_code,
        )


class CSSVariableResolver:
    """Resolves CSS custom properties (variables)."""

    def __init__(self, css_text: str):
        self.variables: Dict[str, str] = {}
        self._parse_root_variables(css_text)

    def _parse_root_variables(self, css: str):
        root_match = re.search(r":root\s*\{([^}]+)\}", css)
        if not root_match:
            return
        for line in root_match.group(1).split(";"):
            line = line.strip()
            if ":" in line and line.startswith("--"):
                name, _, value = line.partition(":")
                self.variables[name.strip()] = value.strip()

    def resolve(self, value: str) -> str:
        """Resolve var(--name) or var(--name, fallback)."""
        def replacer(m):
            inner = m.group(1).strip()
            parts = inner.split(",", 1)
            var_name = parts[0].strip()
            fallback = parts[1].strip() if len(parts) > 1 else ""
            return self.variables.get(var_name, fallback)
        return re.sub(r"var\(([^)]+)\)", replacer, value)

    @staticmethod
    def parse_color(value: str) -> Optional[str]:
        """Parse a CSS color to 6-digit hex (no #). Handles #hex, rgb()."""
        value = value.strip()
        if value.startswith("#"):
            hex_val = value[1:]
            if len(hex_val) == 3:
                hex_val = "".join(c * 2 for c in hex_val)
            if len(hex_val) == 6:
                return hex_val.upper()
        rgb_match = re.match(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)", value)
        if rgb_match:
            r, g, b = int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3))
            return f"{r:02X}{g:02X}{b:02X}"
        return None


def _parse_inline_style(style_str: str) -> Dict[str, str]:
    """Parse an inline style attribute into a dict."""
    props = {}
    if not style_str:
        return props
    for part in style_str.split(";"):
        part = part.strip()
        if ":" in part:
            key, _, val = part.partition(":")
            props[key.strip().lower()] = val.strip()
    return props


class HTMLToDocx:
    """Convert HTML to Word document."""

    def __init__(self):
        self.doc = Document()
        self.css = CSSVariableResolver("")
        self.all_ids: Set[str] = set()
        self.stats = {"headings": 0, "paragraphs": 0, "tables": 0, "lists": 0}

    def convert(self, html: str, output_path: str) -> str:
        """Convert HTML string to DOCX file. Returns output path."""
        soup = BeautifulSoup(html, "html.parser")

        # Extract CSS variables from <style> blocks
        css_text = "\n".join(s.string or "" for s in soup.find_all("style"))
        self.css = CSSVariableResolver(css_text)

        # Pass 1: collect all element IDs for bookmark generation
        for tag in soup.find_all(True):
            if tag.get("id"):
                self.all_ids.add(tag["id"])

        # Setup document defaults
        self._setup_doc()

        # Pass 2: convert content
        body = soup.find("body") or soup
        self._process_children(body)

        # Save output
        abs_path = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(abs_path) or ".", exist_ok=True)
        self.doc.save(abs_path)
        return abs_path

    def convert_file(self, html_path: str, output_path: str) -> str:
        """Convert HTML file to DOCX. Returns output path."""
        with open(html_path, "r", encoding="utf-8") as f:
            html = f.read()
        return self.convert(html, output_path)

    def _setup_doc(self):
        """Configure document margins and default font."""
        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    # ── Dispatch ──────────────────────────────────────────────

    def _process_children(self, parent: Tag):
        """Walk children and dispatch to the right handler."""
        for child in parent.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    para = self.doc.add_paragraph()
                    run = para.add_run(text)
                    self.stats["paragraphs"] += 1
                continue

            if not isinstance(child, Tag):
                continue

            tag_name = child.name.lower()

            if tag_name in SKIP_TAGS:
                continue
            elif tag_name in HEADING_TAGS:
                self._process_heading(child)
            elif tag_name == "p":
                self._process_paragraph(child)
            elif tag_name == "table":
                self._process_table(child)
            elif tag_name in ("ul", "ol"):
                self._process_list(child, ordered=(tag_name == "ol"))
            elif tag_name == "pre":
                self._process_code_block(child)
            elif tag_name == "blockquote":
                self._process_blockquote(child)
            elif tag_name == "hr":
                self._process_hr()
            elif tag_name == "img":
                self._process_image(child)
            elif tag_name == "br":
                self.doc.add_paragraph()
            elif tag_name in CONTAINER_TAGS:
                self._process_children(child)
            elif tag_name in INLINE_TAGS:
                # Inline at block level — wrap in a paragraph
                para = self.doc.add_paragraph()
                self._process_inline(child, para, RunFormat())
                self.stats["paragraphs"] += 1
            else:
                # Unknown block element — recurse into children
                self._process_children(child)

    # ── Headings ──────────────────────────────────────────────

    def _process_heading(self, tag: Tag):
        """Convert h1-h6 to Word headings with bookmarks."""
        level = HEADING_LEVELS.get(tag.name.lower(), 0)
        text = tag.get_text()
        heading = self.doc.add_heading(text, level=level)

        # Apply font size
        size = HEADING_SIZES.get(tag.name.lower(), 12)
        for run in heading.runs:
            run.font.size = Pt(size)

        # Apply color from inline style
        style = _parse_inline_style(tag.get("style", ""))
        if "color" in style:
            resolved = self.css.resolve(style["color"])
            hex_color = CSSVariableResolver.parse_color(resolved)
            if hex_color:
                for run in heading.runs:
                    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)

        # Add bookmark if element has an ID
        if tag.get("id") and tag["id"] in self.all_ids:
            self._add_bookmark(heading, tag["id"])

        self.stats["headings"] += 1

    # ── Paragraphs ────────────────────────────────────────────

    def _process_paragraph(self, tag: Tag):
        """Convert <p> to a Word paragraph with inline formatting."""
        para = self.doc.add_paragraph()
        fmt = RunFormat()

        # Check for paragraph-level color
        style = _parse_inline_style(tag.get("style", ""))
        if "color" in style:
            resolved = self.css.resolve(style["color"])
            hex_color = CSSVariableResolver.parse_color(resolved)
            if hex_color:
                fmt.color = hex_color

        # Check alignment
        if "text-align" in style:
            align = style["text-align"].lower()
            if align == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self._process_inline_children(tag, para, fmt)

        # Add bookmark if element has an ID
        if tag.get("id") and tag["id"] in self.all_ids:
            self._add_bookmark(para, tag["id"])

        self.stats["paragraphs"] += 1

    # ── Tables ────────────────────────────────────────────────

    def _process_table(self, tag: Tag):
        """Convert <table> to Word table with header row formatting."""
        # Collect rows from thead and tbody
        header_rows: List[Tag] = []
        body_rows: List[Tag] = []

        thead = tag.find("thead")
        tbody = tag.find("tbody")

        if thead:
            header_rows = thead.find_all("tr")
        if tbody:
            body_rows = tbody.find_all("tr")

        # If no thead/tbody, all <tr> are body rows
        if not thead and not tbody:
            all_rows = tag.find_all("tr", recursive=False)
            # Check if first row has <th> elements
            if all_rows and all_rows[0].find("th"):
                header_rows = [all_rows[0]]
                body_rows = all_rows[1:]
            else:
                body_rows = all_rows

        all_rows = header_rows + body_rows
        if not all_rows:
            return

        # Determine column count from widest row
        max_cols = 0
        for row in all_rows:
            cells = row.find_all(["td", "th"])
            max_cols = max(max_cols, len(cells))

        if max_cols == 0:
            return

        table = self.doc.add_table(rows=len(all_rows), cols=max_cols)
        table.style = "Table Grid"

        for row_idx, row_tag in enumerate(all_rows):
            cells = row_tag.find_all(["td", "th"])
            is_header = row_idx < len(header_rows)

            for col_idx, cell_tag in enumerate(cells):
                if col_idx >= max_cols:
                    break
                cell = table.cell(row_idx, col_idx)
                cell.text = ""  # Clear default paragraph

                # Add cell content
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                para.text = ""
                fmt = RunFormat(bold=is_header)
                self._process_inline_children(cell_tag, para, fmt)

                # Header row shading
                if is_header:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

                # Cell-level background from style
                cell_style = _parse_inline_style(cell_tag.get("style", ""))
                if "background-color" in cell_style and not is_header:
                    resolved = self.css.resolve(cell_style["background-color"])
                    hex_color = CSSVariableResolver.parse_color(resolved)
                    if hex_color:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
                        cell._tc.get_or_add_tcPr().append(shading)

        self.stats["tables"] += 1

    # ── Lists ─────────────────────────────────────────────────

    def _process_list(self, tag: Tag, ordered: bool = False, level: int = 0):
        """Convert <ul>/<ol> to Word list paragraphs."""
        items = tag.find_all("li", recursive=False)
        for idx, item in enumerate(items):
            # Create paragraph with list style
            if ordered:
                para = self.doc.add_paragraph(style="List Number")
            else:
                para = self.doc.add_paragraph(style="List Bullet")

            # Set indentation for nested lists
            if level > 0:
                para.paragraph_format.left_indent = Cm(1.27 * level)

            # Process inline content (skip nested lists)
            fmt = RunFormat()
            for child in item.children:
                if isinstance(child, NavigableString):
                    text = str(child)
                    if text.strip():
                        run = para.add_run(text)
                        self._apply_run_format(run, fmt)
                elif isinstance(child, Tag):
                    if child.name in ("ul", "ol"):
                        # Nested list — process recursively
                        self._process_list(child, ordered=(child.name == "ol"), level=level + 1)
                    elif child.name in INLINE_TAGS:
                        self._process_inline(child, para, fmt)
                    else:
                        run = para.add_run(child.get_text())
                        self._apply_run_format(run, fmt)

        self.stats["lists"] += 1

    # ── Code Blocks ───────────────────────────────────────────

    def _process_code_block(self, tag: Tag):
        """Convert <pre> (with optional <code>) to monospace paragraph."""
        code_tag = tag.find("code")
        text = (code_tag or tag).get_text()

        para = self.doc.add_paragraph()
        # Light gray background via shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
        para._p.get_or_add_pPr().append(shading)

        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

        self.stats["paragraphs"] += 1

    # ── Blockquotes ───────────────────────────────────────────

    def _process_blockquote(self, tag: Tag):
        """Convert <blockquote> to indented italic paragraph."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.5)

        fmt = RunFormat(italic=True)

        # Check for nested <p> tags
        inner_paras = tag.find_all("p")
        if inner_paras:
            for p_tag in inner_paras:
                self._process_inline_children(p_tag, para, fmt)
        else:
            self._process_inline_children(tag, para, fmt)

        self.stats["paragraphs"] += 1

    # ── Horizontal Rules ──────────────────────────────────────

    def _process_hr(self):
        """Add a horizontal rule as a bordered paragraph."""
        para = self.doc.add_paragraph()
        # Use bottom border to simulate HR
        pPr = para._p.get_or_add_pPr()
        borders = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
            '</w:pBdr>'
        )
        pPr.append(borders)

    # ── Images ────────────────────────────────────────────────

    def _process_image(self, tag: Tag):
        """Convert <img> to inline image if file exists."""
        src = tag.get("src", "")
        if not src or not os.path.isfile(src):
            # Skip missing or remote images
            alt = tag.get("alt", "[image]")
            para = self.doc.add_paragraph()
            run = para.add_run(f"[{alt}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        para = self.doc.add_paragraph()
        run = para.add_run()
        try:
            run.add_picture(src, width=Inches(5.0))
        except Exception:
            run.text = f"[Image: {src}]"

    # ── Inline Processing ─────────────────────────────────────

    def _process_inline_children(self, parent: Tag, para, fmt: RunFormat):
        """Process all children of a block element as inline content."""
        for child in parent.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = para.add_run(text)
                    self._apply_run_format(run, fmt)
            elif isinstance(child, Tag):
                if child.name in SKIP_TAGS:
                    continue
                elif child.name == "br":
                    run = para.add_run()
                    run.add_break()
                elif child.name in INLINE_TAGS:
                    self._process_inline(child, para, fmt)
                elif child.name == "img":
                    alt = child.get("alt", "[image]")
                    run = para.add_run(f"[{alt}]")
                    run.font.italic = True
                else:
                    # Treat unknown inline as plain text
                    run = para.add_run(child.get_text())
                    self._apply_run_format(run, fmt)

    def _process_inline(self, tag: Tag, para, fmt: RunFormat):
        """Recursively process an inline element, accumulating formatting."""
        new_fmt = fmt.copy()
        tag_name = tag.name.lower()

        # Apply formatting based on tag type
        if tag_name in ("strong", "b"):
            new_fmt.bold = True
        elif tag_name in ("em", "i"):
            new_fmt.italic = True
        elif tag_name == "u":
            new_fmt.underline = True
        elif tag_name == "code":
            new_fmt.is_code = True
        elif tag_name == "mark":
            new_fmt.bg_color = "FFFF00"
        elif tag_name == "s":
            pass  # Strikethrough not in RunFormat — skip for simplicity

        # Apply inline style overrides
        style = _parse_inline_style(tag.get("style", ""))
        if "color" in style:
            resolved = self.css.resolve(style["color"])
            hex_color = CSSVariableResolver.parse_color(resolved)
            if hex_color:
                new_fmt.color = hex_color
        if "background-color" in style:
            resolved = self.css.resolve(style["background-color"])
            hex_color = CSSVariableResolver.parse_color(resolved)
            if hex_color:
                new_fmt.bg_color = hex_color
        if "font-weight" in style:
            weight = style["font-weight"].lower()
            if weight in ("bold", "700", "800", "900"):
                new_fmt.bold = True
        if "font-style" in style:
            if style["font-style"].lower() == "italic":
                new_fmt.italic = True

        # Recursively process children
        for child in tag.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = para.add_run(text)
                    self._apply_run_format(run, new_fmt)
            elif isinstance(child, Tag):
                if child.name == "br":
                    run = para.add_run()
                    run.add_break()
                elif child.name in INLINE_TAGS:
                    self._process_inline(child, para, new_fmt)
                else:
                    run = para.add_run(child.get_text())
                    self._apply_run_format(run, new_fmt)

    # ── Run Formatting ────────────────────────────────────────

    def _apply_run_format(self, run, fmt: RunFormat):
        """Apply accumulated formatting to a Word run."""
        if fmt.bold:
            run.font.bold = True
        if fmt.italic:
            run.font.italic = True
        if fmt.underline:
            run.font.underline = True
        if fmt.color:
            hex_color = fmt.color
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        if fmt.bg_color:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fmt.bg_color}" w:val="clear"/>')
            run._r.get_or_add_rPr().append(shading)
        if fmt.is_code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif fmt.font_size:
            run.font.size = Pt(fmt.font_size)

    # ── Bookmarks ─────────────────────────────────────────────

    def _add_bookmark(self, paragraph, bookmark_id: str):
        """Add a Word bookmark to a paragraph element."""
        # Sanitize bookmark name (Word requires alphanumeric + underscore)
        safe_name = re.sub(r"[^a-zA-Z0-9_]", "_", bookmark_id)
        if not safe_name:
            return

        # Generate a numeric ID for the bookmark
        bm_id = str(abs(hash(safe_name)) % 100000)

        start = parse_xml(
            f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="{safe_name}"/>'
        )
        end = parse_xml(
            f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>'
        )

        p_element = paragraph._p if hasattr(paragraph, '_p') else paragraph._element
        p_element.insert(0, start)
        p_element.append(end)
