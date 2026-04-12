"""Tests for html2docx converter."""

import os
import tempfile
import pytest
from docx import Document

from html2docx import HTMLToDocx, CSSVariableResolver, RunFormat


@pytest.fixture
def converter():
    return HTMLToDocx()


@pytest.fixture
def tmp_dir():
    with tempfile.TemporaryDirectory() as d:
        yield d


# ── 1. Basic paragraph conversion ────────────────────────────

def test_basic_paragraph(converter, tmp_dir):
    html = "<p>Hello, world!</p>"
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    assert "Hello, world!" in texts
    assert converter.stats["paragraphs"] >= 1


# ── 2. Heading levels (h1-h6) ────────────────────────────────

def test_heading_levels(converter, tmp_dir):
    html = "".join(f"<h{i}>Heading {i}</h{i}>" for i in range(1, 7))
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    assert converter.stats["headings"] == 6
    # Check that headings appear in output
    all_text = " ".join(p.text for p in doc.paragraphs)
    for i in range(1, 7):
        assert f"Heading {i}" in all_text


# ── 3. Bold/italic inline formatting ─────────────────────────

def test_bold_italic(converter, tmp_dir):
    html = "<p><strong>Bold</strong> and <em>italic</em> text</p>"
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    # Find the paragraph with our content
    para = [p for p in doc.paragraphs if "Bold" in p.text][0]
    runs = para.runs
    bold_runs = [r for r in runs if r.font.bold]
    italic_runs = [r for r in runs if r.font.italic]
    assert len(bold_runs) >= 1
    assert len(italic_runs) >= 1


# ── 4. Table with header row ─────────────────────────────────

def test_table_with_header(converter, tmp_dir):
    html = """
    <table>
      <thead><tr><th>Name</th><th>Value</th></tr></thead>
      <tbody><tr><td>Alpha</td><td>100</td></tr></tbody>
    </table>
    """
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Name"
    assert table.cell(0, 1).text == "Value"
    assert table.cell(1, 0).text == "Alpha"
    assert table.cell(1, 1).text == "100"
    assert converter.stats["tables"] == 1


# ── 5. Ordered and unordered lists ───────────────────────────

def test_unordered_list(converter, tmp_dir):
    html = "<ul><li>Item A</li><li>Item B</li><li>Item C</li></ul>"
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    list_texts = [p.text for p in doc.paragraphs if "Item" in p.text]
    assert len(list_texts) == 3
    assert converter.stats["lists"] >= 1


def test_ordered_list(converter, tmp_dir):
    html = "<ol><li>First</li><li>Second</li></ol>"
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    list_texts = [p.text for p in doc.paragraphs if "First" in p.text or "Second" in p.text]
    assert len(list_texts) == 2


# ── 6. CSS variable resolution ───────────────────────────────

def test_css_variable_resolution():
    css = ":root { --primary: #1a2744; --accent: #ff6600; }"
    resolver = CSSVariableResolver(css)
    assert resolver.resolve("var(--primary)") == "#1a2744"
    assert resolver.resolve("var(--accent)") == "#ff6600"
    assert resolver.resolve("var(--missing, blue)") == "blue"
    assert resolver.resolve("var(--missing)") == ""


def test_css_variable_in_html(converter, tmp_dir):
    html = """
    <style>:root { --heading-color: #336699; }</style>
    <h1 style="color: var(--heading-color)">Styled Heading</h1>
    """
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    heading = [p for p in doc.paragraphs if "Styled Heading" in p.text][0]
    # The heading should have a colored run
    colored_runs = [r for r in heading.runs if r.font.color.rgb is not None]
    assert len(colored_runs) >= 1


# ── 7. Color parsing (hex, rgb) ──────────────────────────────

def test_parse_color_hex():
    assert CSSVariableResolver.parse_color("#1a2744") == "1A2744"
    assert CSSVariableResolver.parse_color("#abc") == "AABBCC"
    assert CSSVariableResolver.parse_color("#ABC") == "AABBCC"


def test_parse_color_rgb():
    assert CSSVariableResolver.parse_color("rgb(255, 0, 128)") == "FF0080"
    assert CSSVariableResolver.parse_color("rgb(0, 0, 0)") == "000000"


def test_parse_color_invalid():
    assert CSSVariableResolver.parse_color("not-a-color") is None
    assert CSSVariableResolver.parse_color("") is None


# ── 8. File input/output ─────────────────────────────────────

def test_file_conversion(converter, tmp_dir):
    html_path = os.path.join(tmp_dir, "input.html")
    docx_path = os.path.join(tmp_dir, "output.docx")

    with open(html_path, "w", encoding="utf-8") as f:
        f.write("<html><body><p>File test</p></body></html>")

    out = converter.convert_file(html_path, docx_path)
    assert os.path.isfile(out)
    doc = Document(out)
    assert any("File test" in p.text for p in doc.paragraphs)


# ── 9. Nested inline elements ────────────────────────────────

def test_nested_inline(converter, tmp_dir):
    html = "<p><em><strong>Bold inside italic</strong></em></p>"
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    para = [p for p in doc.paragraphs if "Bold inside italic" in p.text][0]
    # Should have a run that is both bold and italic
    both = [r for r in para.runs if r.font.bold and r.font.italic]
    assert len(both) >= 1


# ── 10. Code blocks ──────────────────────────────────────────

def test_code_block(converter, tmp_dir):
    html = "<pre><code>print('hello')</code></pre>"
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    code_paras = [p for p in doc.paragraphs if "print" in p.text]
    assert len(code_paras) >= 1
    run = code_paras[0].runs[0]
    assert run.font.name == "Consolas"
    assert run.font.size == 114300  # Pt(9) in EMU


# ── Bonus: RunFormat dataclass ────────────────────────────────

def test_runformat_copy():
    fmt = RunFormat(bold=True, color="FF0000", is_code=True)
    copy = fmt.copy()
    assert copy.bold is True
    assert copy.color == "FF0000"
    assert copy.is_code is True
    # Modifying copy doesn't affect original
    copy.bold = False
    assert fmt.bold is True


# ── Bonus: Bookmarks ─────────────────────────────────────────

def test_bookmark_generation(converter, tmp_dir):
    html = '<h1 id="intro">Introduction</h1><p>Content here</p>'
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    assert os.path.isfile(out)
    assert "intro" in converter.all_ids


# ── Bonus: Multiple paragraphs ───────────────────────────────

def test_multiple_paragraphs(converter, tmp_dir):
    html = "<p>First</p><p>Second</p><p>Third</p>"
    out = converter.convert(html, os.path.join(tmp_dir, "out.docx"))
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    assert "First" in texts
    assert "Second" in texts
    assert "Third" in texts
